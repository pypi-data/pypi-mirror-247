import openpyxl
import docx
import shutil
import os


def FileName(FilePath):
    return os.path.basename(FilePath)


def FileFatherPath(FilePath):
    return os.path.dirname(FilePath)


def CopyFile(FileOriginalPath, FileNewPath):
    shutil.copy2(FileOriginalPath, FileNewPath)


def MoveFile(FileOriginalPath, FileNewPath):
    shutil.move(FileOriginalPath, FileNewPath)


def ExcleRead(ExclePath, SheetIndex, Rowlow, Rowmax, Collow, Colmax):
    if Rowlow > Rowmax:
        t = Rowmax
        Rowmax = Rowlow
        Rowlow = t
    if Collow > Colmax:
        t = Rowmax
        Colmax = Collow
        Collow = t
    RowNum = Rowmax - Rowlow + 1
    ColNum = Colmax - Collow + 1
    # 打开工作簿
    workbook = openpyxl.load_workbook(ExclePath)
    # 获取所有工作表
    sheets = workbook.sheetnames
    # 选择第一个工作表
    sheet = workbook[sheets[SheetIndex - 1]]
    # 存储为列表

    m = 0
    i = 0
    SheetData = [[None for j in range(Colmax - Collow + 1)] for i in range(Rowmax - Rowlow + 1)]
    for row in sheet.iter_rows():
        n = 0
        j = 0
        for cell in row:

            if m + 1 >= Rowlow and m + 1 <= Rowmax and n + 1 >= Collow and n + 1 <= Colmax:
                # 获取单元格的值
                a = cell.value
                SheetData[i][j] = a
                j = j + 1
            n = n + 1
        if m + 1 >= Rowlow and m + 1 <= Rowmax:
            i = i + 1
        m = m + 1
    if RowNum == 1:
        SheetSingle = {}
        i = 0
        for data in SheetData[0]:
            SheetSingle[i] = data
            i = i + 1
        SheetData = SheetSingle
    elif ColNum == 1:
        SheetSingle = {}
        for i in range(0, RowNum):
            SheetSingle[i] = SheetData[i][0]
        SheetData = SheetSingle
    return SheetData


def ExcleWrite(ExclePath, SheetIndex, Row, Col, Value, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(ExclePath)
        if FileFatherPath(ExclePath) != "":
            ExcleNewPath = FileFatherPath(ExclePath) + "\\" + FileNewName
        else:
            ExcleNewPath = FileNewName
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExcleNewPath)
    else:
        workbook = openpyxl.load_workbook(ExclePath)
        sheet_names = workbook.sheetnames
        SheetName = sheet_names[SheetIndex - 1]
        # 选择要操作的工作表
        sheet = workbook[SheetName]
        # 在指定的单元格写入数据
        sheet.cell(row=Row, column=Col, value=Value)
        # 保存文件
        workbook.save(ExclePath)


def WordTableRead(WordPath, TableIndex):
    doc = docx.Document(WordPath)
    table = doc.tables[TableIndex - 1]
    RowNum = 0
    for row in table.rows:
        ColNum = 0
        for cell in row.cells:
            ColNum = ColNum + 1
        RowNum = RowNum + 1
    SheetData = [[None for j in range(ColNum)] for i in range(RowNum)]
    i = 0

    for row in table.rows:
        j = 0
        for cell in row.cells:
            if i == 0 and j == 0:
                bcell_text = None
                cell_text = cell.text
                SheetData[i][j] = cell_text
            else:
                bcell_text = cell_text
                cell_text = cell.text
                if bcell_text != cell_text:
                    SheetData[i][j] = cell_text
                else:
                    SheetData[i][j] = None
            j = j + 1
        i = i + 1
    return SheetData


def WordTableWrite(WordPath, TableIndex, Row, Col, Text, SaveAsNewFile):
    if SaveAsNewFile:
        FileNewName = "New_" + FileName(WordPath)
        if FileFatherPath(WordPath) != "":
            WordNewPath = FileFatherPath(WordPath) + "\\" + FileNewName
        else:
            WordNewPath = FileNewName
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        Cell.text = Text
        doc.save(WordNewPath)
    else:
        doc = docx.Document(WordPath)
        table = doc.tables[TableIndex - 1]
        Cell = table.cell(Row - 1, Col - 1)
        Cell.text = Text
        doc.save(WordPath)
